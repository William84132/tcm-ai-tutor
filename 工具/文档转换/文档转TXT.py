# -*- coding: utf-8 -*-
"""
文档转TXT v1.0 — 本地一键转换 PDF / DOCX / DOC → TXT
====================================================
零 token、零联网、纯本地运行。拖放文件或文件夹到 文档转TXT.bat 上即可。

支持格式与策略:
  PDF  文本型 → pdfplumber 直接提取（秒级，不 OCR）
       扫描型 → poppler pdftoppm 渲染 300dpi + tesseract OCR（chi_sim+chi_tra）
       （自动判定：平均每页有效字符 < 15 个即视为扫描版）
  DOCX → python-docx 提取段落与表格
  DOC  → Word COM（本机安装 Office 时可用）另存为文本

输出: 每个源文件所在目录的 转换输出/ 子目录下，同名 .txt

工具路径配置: 读取 <仓库根>/叶天士/08-学习工具/工具配置.json
（tesseract / tessdata / poppler_bin 三个字段；不配置则尝试 PATH 自动查找）
"""
import os
import sys
import json
import glob
import shutil
import subprocess
import tempfile

sys.stdout.reconfigure(encoding='utf-8')

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
REPO_ROOT = os.path.dirname(os.path.dirname(SCRIPT_DIR))  # 工具/文档转换 -> 仓库根

LIMIT_PAGES = None          # 测试用：仅处理前 N 页（None = 全部）
MIN_CHARS_PER_PAGE = 15     # 平均每页有效字符低于此值 -> 判定扫描版转 OCR
OCR_DPI = 300               # OCR 渲染分辨率


def log(msg):
    print(msg, flush=True)


def load_cfg():
    cfg = {}
    p = os.path.join(REPO_ROOT, '叶天士', '08-学习工具', '工具配置.json')
    if os.path.isfile(p):
        try:
            with open(p, encoding='utf-8') as f:
                cfg = json.load(f)
        except Exception as e:
            log('[警告] 工具配置.json 读取失败: %s' % e)
    else:
        log('[提示] 未找到 工具配置.json（%s），将尝试从 PATH 查找工具' % p)
    return cfg


def find_tool(cfg, key, names):
    """返回工具可执行文件路径；cfg 优先，其次 PATH"""
    v = cfg.get(key, '')
    if v and os.path.isfile(v):
        return v
    for n in names:
        w = shutil.which(n)
        if w:
            return w
    return None


CFG = load_cfg()
TESSERACT = find_tool(CFG, 'tesseract', ['tesseract.exe', 'tesseract'])
TESSDATA = CFG.get('tessdata', '')
POPPLER_BIN = CFG.get('poppler_bin', '')
POPPLER_PPPM = os.path.join(POPPLER_BIN, 'pdftoppm.exe') if POPPLER_BIN else shutil.which('pdftoppm')


def run_silent(args):
    """运行外部程序，不捕获输出（避免管道限制），仅返回退出码"""
    try:
        return subprocess.run(args, stdout=subprocess.DEVNULL,
                              stderr=subprocess.DEVNULL).returncode
    except Exception as e:
        log('  [失败] 无法运行 %s: %s' % (args[0], e))
        return -1


# ==================== PDF ====================

def extract_pdf_text(path):
    """文本型 PDF 提取，返回 (文本, 页数, 平均每页字符数)"""
    import pdfplumber
    parts, total, pages = [], 0, 0
    with pdfplumber.open(path) as pdf:
        for i, page in enumerate(pdf.pages):
            if LIMIT_PAGES and i >= LIMIT_PAGES:
                break
            t = page.extract_text() or ''
            parts.append(t)
            total += len(t.strip())
            pages += 1
    return '\n'.join(parts), pages, (total / pages if pages else 0)


def ocr_pdf(path):
    """扫描版 PDF OCR：pypdfium2 渲染 300dpi -> tesseract 识别，返回文本"""
    if not TESSERACT:
        return None, '未找到 tesseract（请检查 工具配置.json 的 tesseract/tessdata 字段）'
    if TESSDATA and os.path.isdir(TESSDATA):
        os.environ['TESSDATA_PREFIX'] = TESSDATA
    try:
        import pypdfium2 as pdfium
        from PIL import Image
    except ImportError:
        return None, '未安装 pypdfium2 / pillow（pip install pypdfium2 pillow）'
    with tempfile.TemporaryDirectory() as td:
        try:
            pdf = pdfium.PdfDocument(path)
        except Exception as e:
            return None, 'PDF 打开失败: %s' % e
        n = len(pdf)
        pages = range(min(n, LIMIT_PAGES)) if LIMIT_PAGES else range(n)
        parts = []
        for i in pages:
            page = pdf[i]
            bmp = page.render(scale=300 / 72.0)  # 300 dpi
            img = bmp.to_pil()
            png = os.path.join(td, 'pg%03d.png' % (i + 1))
            img.save(png)
            out_base = os.path.join(td, 'out%03d' % (i + 1))
            rc = run_silent([TESSERACT, png, out_base, '-l', 'chi_sim+chi_tra'])
            if rc != 0:
                return None, 'tesseract OCR 失败（第 %d 页）' % (i + 1)
            txt = out_base + '.txt'
            if os.path.isfile(txt):
                with open(txt, encoding='utf-8', errors='replace') as f:
                    t = f.read()
                if t.strip():          # 跳过无可识别文字的页（如图片封面）
                    parts.append(t)
        pdf.close()
        if not parts:
            return None, '未产出任何 OCR 文本'
        return '\n'.join(parts), None


def convert_pdf(path, out_dir):
    os.makedirs(out_dir, exist_ok=True)
    name = os.path.splitext(os.path.basename(path))[0]
    out_path = os.path.join(out_dir, name + '.txt')
    text, pages, per_page = extract_pdf_text(path)
    if per_page and per_page >= MIN_CHARS_PER_PAGE:
        mode = '文本提取（%d 页，平均每页 %d 字符）' % (pages, per_page)
        with open(out_path, 'w', encoding='utf-8') as f:
            f.write(text)
        return True, mode
    log('  [自动] 判定为扫描版（平均每页 %s 字符 < %d），启动 OCR ...'
        % (('%.0f' % per_page) if per_page else '0', MIN_CHARS_PER_PAGE))
    text, err = ocr_pdf(path)
    if text is None:
        return False, 'OCR 失败: %s' % err
    with open(out_path, 'w', encoding='utf-8') as f:
        f.write(text)
    return True, 'OCR 识别完成'


# ==================== DOCX / DOC ====================

def convert_docx(path, out_dir):
    os.makedirs(out_dir, exist_ok=True)
    import docx
    name = os.path.splitext(os.path.basename(path))[0]
    d = docx.Document(path)
    parts = []
    for p in d.paragraphs:
        if p.text.strip():
            parts.append(p.text)
    for tbl in d.tables:
        for row in tbl.rows:
            cells = [c.text.strip() for c in row.cells]
            parts.append(' | '.join(cells))
    out_path = os.path.join(out_dir, name + '.txt')
    with open(out_path, 'w', encoding='utf-8') as f:
        f.write('\n'.join(parts))
    return True, 'DOCX 提取（%d 段）' % len(parts)


def convert_doc(path, out_dir):
    """旧版 .doc：Word COM 另存为 Unicode 文本"""
    os.makedirs(out_dir, exist_ok=True)
    try:
        import win32com.client
    except ImportError:
        return False, '未安装 pywin32，无法转换 .doc；请用 Word 另存为 .docx 后重试'
    name = os.path.splitext(os.path.basename(path))[0]
    out_path = os.path.join(out_dir, name + '.txt')
    word = None
    try:
        word = win32com.client.Dispatch('Word.Application')
        word.Visible = False
        doc = word.Documents.Open(os.path.abspath(path), ReadOnly=True)
        doc.SaveAs2(os.path.abspath(out_path), FileFormat=7)  # wdFormatUnicodeText
        doc.Close(False)
        # Word 纯文本可能按系统 ANSI 编码输出，统一归一为 UTF-8
        with open(out_path, 'rb') as f:
            raw = f.read()
        try:
            txt = raw.decode('utf-8')
        except UnicodeDecodeError:
            txt = raw.decode('gb18030', errors='replace')
        with open(out_path, 'w', encoding='utf-8') as f:
            f.write(txt)
        return True, 'DOC 转换完成（Word COM）'
    except Exception as e:
        return False, 'Word COM 失败: %s' % e
    finally:
        if word:
            try:
                word.Quit()
            except Exception:
                pass


# ==================== 主流程 ====================

def collect_inputs(paths):
    """展开输入：目录递归收集支持的文件"""
    exts = ('.pdf', '.docx', '.doc')
    files = []
    for p in paths:
        p = os.path.abspath(p)
        if os.path.isdir(p):
            for root, dirs, fs in os.walk(p):
                dirs[:] = [d for d in dirs if d not in ('转换输出',)]
                for fn in fs:
                    if fn.lower().endswith(exts):
                        files.append(os.path.join(root, fn))
        elif os.path.isfile(p) and p.lower().endswith(exts):
            files.append(p)
    return sorted(set(files))


def main():
    if len(sys.argv) < 2:
        print('=' * 60)
        print(' 文档转TXT v1.0 — 本地转换 PDF/DOCX/DOC → TXT')
        print('=' * 60)
        print('用法: 把 PDF/DOCX/DOC 文件或文件夹拖到 文档转TXT.bat 上')
        print('      或命令行: python 文档转TXT.py <文件或文件夹> [更多...]')
        print('输出: 每个源文件旁的 转换输出/ 目录（同名 .txt）')
        return 1
    files = collect_inputs(sys.argv[1:])
    if not files:
        print('[X] 未找到可转换的文件（支持 .pdf / .docx / .doc）')
        return 1
    print('=' * 60)
    print(' 文档转TXT — 共 %d 个文件' % len(files))
    print('=' * 60)
    ok = fail = 0
    for i, p in enumerate(files, 1):
        ext = os.path.splitext(p)[1].lower()
        out_dir = os.path.join(os.path.dirname(p), '转换输出')
        os.makedirs(out_dir, exist_ok=True)
        log('[%d/%d] %s' % (i, len(files), os.path.basename(p)))
        try:
            if ext == '.pdf':
                okk, mode = convert_pdf(p, out_dir)
            elif ext == '.docx':
                okk, mode = convert_docx(p, out_dir)
            else:
                okk, mode = convert_doc(p, out_dir)
        except Exception as e:
            okk, mode = False, '异常: %s' % e
        if okk:
            ok += 1
            log('  [OK] %s -> 转换输出/%s.txt' % (mode, os.path.splitext(os.path.basename(p))[0]))
        else:
            fail += 1
            log('  [X] %s' % mode)
    print('=' * 60)
    print('完成: 成功 %d，失败 %d' % (ok, fail))
    return 0 if fail == 0 else 1


if __name__ == '__main__':
    sys.exit(main())
